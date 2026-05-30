from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path('reports/fujimoto_126_strategy_report_2026-05-30.md')
OUT = Path('reports/documents/fujimoto_126_strategy_report_2026-05-30.docx')
OUT.parent.mkdir(parents=True, exist_ok=True)

md = SRC.read_text(encoding='utf-8')

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text.strip())
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(max_cols):
            text = row[j] if j < len(row) else ''
            cell = table.cell(i, j)
            set_cell_text(cell, text, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, 'D9EAF7')
    doc.add_paragraph('')


def flush_table(doc, table_lines):
    if not table_lines:
        return
    rows = []
    for line in table_lines:
        parts = [p.strip() for p in line.strip().strip('|').split('|')]
        if all(re.fullmatch(r':?-{3,}:?', p or '') for p in parts):
            continue
        rows.append(parts)
    add_table(doc, rows)


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    for line in code.rstrip('\n').splitlines():
        run = p.add_run(line + '\n')
        run.font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    doc.add_paragraph('')


def add_paragraph_with_inline(doc, text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    # Basic bold inline handling for **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith('**') and part.endswith('**')
        clean = part[2:-2] if bold else part
        clean = clean.replace('`', '')
        run = p.add_run(clean)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        run.font.size = Pt(10.5)
        run.bold = bold
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = doc.styles
styles['Normal'].font.name = 'Malgun Gothic'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
styles['Normal'].font.size = Pt(10.5)
for style_name, size, color in [
    ('Title', 24, RGBColor(0x1A, 0x36, 0x5D)),
    ('Heading 1', 17, RGBColor(0x1A, 0x36, 0x5D)),
    ('Heading 2', 13, RGBColor(0x2B, 0x6C, 0xB0)),
    ('Heading 3', 11, RGBColor(0x2D, 0x37, 0x48)),
]:
    st = styles[style_name]
    st.font.name = 'Malgun Gothic'
    st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True

lines = md.splitlines()
in_code = False
code_lines = []
table_lines = []
first_heading = True

for raw in lines:
    line = raw.rstrip()
    if line.startswith('```'):
        if in_code:
            add_code_block(doc, '\n'.join(code_lines))
            code_lines = []
            in_code = False
        else:
            flush_table(doc, table_lines)
            table_lines = []
            in_code = True
        continue
    if in_code:
        code_lines.append(line)
        continue
    if line.startswith('|') and line.endswith('|'):
        table_lines.append(line)
        continue
    else:
        flush_table(doc, table_lines)
        table_lines = []
    if not line.strip():
        continue
    if line.strip() == '---':
        doc.add_paragraph('')
        continue
    if line.startswith('# '):
        title = line[2:].strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Malgun Gothic'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        doc.add_paragraph('')
        first_heading = False
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=2)
    elif line.startswith('- '):
        add_paragraph_with_inline(doc, line[2:].strip(), style='List Bullet')
    elif re.match(r'^\d+\.\s+', line):
        add_paragraph_with_inline(doc, re.sub(r'^\d+\.\s+', '', line), style='List Number')
    elif line.startswith('> '):
        p = add_paragraph_with_inline(doc, line[2:].strip())
        p.paragraph_format.left_indent = Inches(0.3)
    else:
        add_paragraph_with_inline(doc, line)

flush_table(doc, table_lines)
if in_code and code_lines:
    add_code_block(doc, '\n'.join(code_lines))

# Footer
for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Dante Labs / Trading Strategy Report / Confidential')
    run.font.name = 'Malgun Gothic'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(OUT)
print(str(OUT))
